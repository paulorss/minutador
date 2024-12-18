import streamlit as st
import os
import re
from docx import Document
import html
from datetime import datetime

class GeradorMinutas:
    def __init__(self):
        # Usa caminho relativo para a pasta modelos
        self.PASTA_MODELOS = os.path.join(os.path.dirname(__file__), 'modelos')
        
    def extrair_campos_e_texto(self, caminho_documento):
        """
        Extrai o texto do documento preservando a formatação e destacando campos
        """
        documento = Document(caminho_documento)
        texto_html = []
        campos = []
        
        # Coleta campos únicos
        todos_campos = []
        for paragrafo in documento.paragraphs:
            campos_paragrafo = re.findall(r'\[(.*?)\]', paragrafo.text)
            todos_campos.extend(campos_paragrafo)
        campos_unicos = list(dict.fromkeys(todos_campos))

        for paragrafo in documento.paragraphs:
            # Alinhamento do parágrafo
            alinhamento = 'left'
            if paragrafo.alignment is not None:
                if paragrafo.alignment == 1:
                    alinhamento = 'center'
                elif paragrafo.alignment == 2:
                    alinhamento = 'right'
                elif paragrafo.alignment == 3:
                    alinhamento = 'justify'
            
            paragrafo_html = f'<p style="text-align: {alinhamento};">'
            
            # Obtém o texto completo do parágrafo
            texto_completo = paragrafo.text
            texto_processado = ""
            ultimo_indice = 0
            
            # Encontra todos os campos no parágrafo
            for match in re.finditer(r'\[(.*?)\]', texto_completo):
                # Adiciona o texto antes do campo
                texto_antes = texto_completo[ultimo_indice:match.start()]
                if texto_antes:
                    # Procura a formatação correta para este trecho
                    for run in paragrafo.runs:
                        if texto_antes in run.text:
                            estilo = []
                            if run.bold:
                                estilo.append('font-weight: bold')
                            if run.italic:
                                estilo.append('font-style: italic')
                            if run.underline:
                                estilo.append('text-decoration: underline')
                            if estilo:
                                texto_antes = f'<span style="{"; ".join(estilo)}">{html.escape(texto_antes)}</span>'
                            break
                    texto_processado += texto_antes
                
                # Adiciona o campo com destaque amarelo
                campo = match.group(0)
                texto_processado += f'<span style="background-color: #FFFF00; padding: 0 2px;">{html.escape(campo)}</span>'
                
                ultimo_indice = match.end()
            
            # Adiciona o texto restante após o último campo
            if ultimo_indice < len(texto_completo):
                texto_restante = texto_completo[ultimo_indice:]
                # Procura a formatação correta para este trecho
                for run in paragrafo.runs:
                    if texto_restante in run.text:
                        estilo = []
                        if run.bold:
                            estilo.append('font-weight: bold')
                        if run.italic:
                            estilo.append('font-style: italic')
                        if run.underline:
                            estilo.append('text-decoration: underline')
                        if estilo:
                            texto_restante = f'<span style="{"; ".join(estilo)}">{html.escape(texto_restante)}</span>'
                        break
                texto_processado += texto_restante
            
            paragrafo_html += texto_processado
            paragrafo_html += '</p>'
            texto_html.append(paragrafo_html)
        
        return {
            'texto_original': '\n'.join(texto_html),
            'campos': campos_unicos
        }

    def gerar_nova_minuta(self, caminho_modelo, dados_campos):
        """
        Gera uma nova minuta preservando exatamente a formatação original
        """
        documento_original = Document(caminho_modelo)
        novo_documento = Document()
        
        for paragrafo_original in documento_original.paragraphs:
            novo_paragrafo = novo_documento.add_paragraph()
            
            # Copia formatação do parágrafo
            if paragrafo_original.alignment is not None:
                novo_paragrafo.alignment = paragrafo_original.alignment
            
            if hasattr(paragrafo_original, 'paragraph_format'):
                try:
                    if paragrafo_original.paragraph_format.first_line_indent:
                        novo_paragrafo.paragraph_format.first_line_indent = paragrafo_original.paragraph_format.first_line_indent
                    if paragrafo_original.paragraph_format.space_before:
                        novo_paragrafo.paragraph_format.space_before = paragrafo_original.paragraph_format.space_before
                    if paragrafo_original.paragraph_format.space_after:
                        novo_paragrafo.paragraph_format.space_after = paragrafo_original.paragraph_format.space_after
                except:
                    pass

            # Obtém o texto do parágrafo e identifica campos
            texto_paragrafo = ''
            runs_info = []
            
            # Coleta informações sobre runs e texto
            for run in paragrafo_original.runs:
                inicio = len(texto_paragrafo)
                texto_paragrafo += run.text
                fim = len(texto_paragrafo)
                runs_info.append({
                    'inicio': inicio,
                    'fim': fim,
                    'run': run,
                    'texto': run.text
                })
                
            # Encontra todos os campos no texto completo
            matches = list(re.finditer(r'\[(.*?)\]', texto_paragrafo))
            
            if matches:
                ultimo_indice = 0
                for match in matches:
                    # Processa texto antes do campo
                    if match.start() > ultimo_indice:
                        texto_antes = texto_paragrafo[ultimo_indice:match.start()]
                        posicao = ultimo_indice
                        
                        # Divide o texto antes em runs preservando formatação original
                        while posicao < match.start():
                            for info in runs_info:
                                if info['inicio'] <= posicao < info['fim']:
                                    # Calcula quanto texto deste run usar
                                    fim_texto = min(match.start() - posicao, info['fim'] - posicao)
                                    texto_run = texto_paragrafo[posicao:posicao + fim_texto]
                                    novo_run = novo_paragrafo.add_run(texto_run)
                                    self._copiar_formatacao_run(novo_run, info['run'])
                                    posicao += len(texto_run)
                                    break
                    
                    # Processa o campo
                    nome_campo = match.group(1).strip()
                    if nome_campo in dados_campos:
                        # Encontra o run que contém o campo original
                        for info in runs_info:
                            if info['inicio'] <= match.start() < info['fim']:
                                novo_run = novo_paragrafo.add_run(dados_campos[nome_campo])
                                self._copiar_formatacao_run(novo_run, info['run'])
                                break
                        else:
                            # Se não encontrar run específico, adiciona sem formatação
                            novo_run = novo_paragrafo.add_run(dados_campos[nome_campo])
                    
                    ultimo_indice = match.end()
                
                # Processa texto após o último campo
                if ultimo_indice < len(texto_paragrafo):
                    posicao = ultimo_indice
                    while posicao < len(texto_paragrafo):
                        for info in runs_info:
                            if info['inicio'] <= posicao < info['fim']:
                                fim_texto = min(len(texto_paragrafo) - posicao, info['fim'] - posicao)
                                texto_run = texto_paragrafo[posicao:posicao + fim_texto]
                                novo_run = novo_paragrafo.add_run(texto_run)
                                self._copiar_formatacao_run(novo_run, info['run'])
                                posicao += len(texto_run)
                                break
            else:
                # Se não há campos, copia exatamente os runs originais
                for run in paragrafo_original.runs:
                    novo_run = novo_paragrafo.add_run(run.text)
                    self._copiar_formatacao_run(novo_run, run)
        
        # Gera nome único para o arquivo
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f'minuta_modificada_{timestamp}.docx'
        
        # Salva em pasta temporária do Streamlit
        caminho_saida = os.path.join(os.path.dirname(__file__), 'minutas_geradas', nome_arquivo)
        os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
        
        novo_documento.save(caminho_saida)
        return caminho_saida

    def _copiar_formatacao_run(self, run_destino, run_origem):
        """
        Copia a formatação de um run para outro
        """
        try:
            run_destino.bold = run_origem.bold
            run_destino.italic = run_origem.italic
            run_destino.underline = run_origem.underline
            
            if run_origem.font.name:
                run_destino.font.name = run_origem.font.name
            if hasattr(run_origem.font, 'size') and run_origem.font.size:
                run_destino.font.size = run_origem.font.size
        except Exception as e:
            st.error(f"Erro ao copiar formatação: {e}")
            pass

def main():
    st.set_page_config(
        page_title="Gerador de Minutas",
        page_icon="📄",
        layout="wide"
    )
    
    st.title('Gerador de Minutas Personalizadas')
    
    if 'caminho_minuta' not in st.session_state:
        st.session_state.caminho_minuta = None
    
    # Inicializa o gerador com caminho relativo
    gerador = GeradorMinutas()
    
    # Lista documentos da pasta modelos
    try:
        documentos = [
            f for f in os.listdir(gerador.PASTA_MODELOS) 
            if f.endswith('.docx') and not f.startswith('~$')
        ]
    except Exception as e:
        st.error(f"Erro ao listar modelos: {e}")
        documentos = []
    
    if not documentos:
        st.warning("Nenhum modelo encontrado na pasta 'modelos'. Adicione arquivos .docx à pasta.")
        return
    
    # Seleção do documento
    documento_selecionado = st.selectbox(
        'Escolha um modelo de minuta', 
        documentos
    )
    
    if documento_selecionado:
        caminho_documento = os.path.join(gerador.PASTA_MODELOS, documento_selecionado)
        resultado = gerador.extrair_campos_e_texto(caminho_documento)
        
        # Layout em duas colunas
        col1, col2 = st.columns([2, 1])
        
        # Texto do modelo (esquerda)
        with col1:
            st.subheader('Modelo de Minuta')
            st.markdown(
                resultado['texto_original'],
                unsafe_allow_html=True
            )
        
        # Campos para preenchimento e botões (direita)
        with col2:
            st.subheader('Preencha os Campos')
            
            with st.form('gerar_minuta'):
                # Campos dinâmicos com valores padrão vazios
                dados_campos = {}
                for campo in resultado['campos']:
                    # Remove os colchetes para exibição
                    campo_limpo = campo.strip()
                    dados_campos[campo_limpo] = st.text_input(
                        f'Campo: {campo_limpo}',
                        key=f'campo_{campo_limpo}'
                    )
                
                submit = st.form_submit_button('Gerar Nova Minuta')
                
                if submit:
                    if all(dados_campos.values()):
                        st.session_state.caminho_minuta = gerador.gerar_nova_minuta(
                            caminho_documento,
                            dados_campos
                        )
                        st.success('Minuta gerada com sucesso!')
                    else:
                        st.warning('Preencha todos os campos')
            
            # Botão de download
            if st.session_state.caminho_minuta:
                with open(st.session_state.caminho_minuta, 'rb') as f:
                    st.download_button(
                        label='Baixar Minuta Modificada',
                        data=f.read(),
                        file_name=os.path.basename(st.session_state.caminho_minuta),
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    
    

if __name__ == '__main__':
    main()
